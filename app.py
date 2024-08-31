import tkinter as tk
from tkinter import ttk, messagebox
import json
import os
import docx
import random
from docx.shared import Pt
import win32print
import win32api
from tkinter import simpledialog
import sys
import time
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import inch
from docx2pdf import convert
from pdf2image import convert_from_path
from PIL import Image
import tempfile


def resource_path(relative_path):
    """Returns the absolute path to a resource, works for development and PyInstaller"""
    try:
        # PyInstallerの実行時には、このコードが _MEIPASSを使用する
        base_path = sys._MEIPASS
    except AttributeError:
        # 開発時やPyInstaller以外の環境ではカレントディレクトリを使用
        base_path = os.path.abspath(".")
    return os.path.join(base_path, relative_path)

# 単語帳の読み込み
def load_vocabulary(file_path):
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"指定された単語帳が見つかりません: {file_path}")
    vocabulary = []
    with open(file_path, 'r', encoding='utf-8') as file:
        for line in file:
            parts = line.strip().split('\t')
            if len(parts) == 3:
                no, word, meaning = parts
                vocabulary.append({"No": int(no), "Word": word, "meaning": meaning})
    return vocabulary


# テスト生成
def create_test():
    try:
        selected_book = book_var.get()
        start_no = int(start_no_entry.get())
        end_no = int(end_no_entry.get())
        num_words = int(num_words_entry.get())
        include_answers = include_answers_var.get()
        num_copies = int(num_copies_entry.get())
        use_image_pdf = True
        printer_name = printer_var.get()

        if start_no <= 0 or start_no > end_no:
                raise ValueError("開始Noが終了Noより大きい、または開始Noが0以下です。")

        file_paths = {
            "システム英単語": resource_path("vocabulary_books/system_English_vocabularyBook.txt"),
            "出る順パス単": resource_path("vocabulary_books/deru_jun_pasutan.txt")
        }

        file_path = file_paths.get(selected_book)
        if not file_path:
            raise ValueError("選択した単語帳のファイルパスが見つかりません。")
        
        if start_no <= 0:
            raise ValueError("開始Noは1以上の整数である必要があります。")
        
        if end_no < start_no:
            raise ValueError("開始Noが終了Noよりも大きいです。")
        
        vocabulary = load_vocabulary(file_path)

        max_no = max(word["No"] for word in vocabulary)
        if end_no > max_no:
            raise ValueError(f"この単語帳は最大No{max_no}までです。")
        
        if num_words > (end_no - start_no + 1):
            raise ValueError(f"テスト問題数がテスト範囲より多いです。")
        
        words_in_range = [word for word in vocabulary if start_no <= word["No"] <= end_no]
        selected_words = random.sample(words_in_range, num_words)
        
        docx_path = create_word_file(selected_words, start_no, end_no)
        
        if use_image_pdf:
            pdf_path = convert_docx_to_image_pdf(docx_path)
        else:
            pdf_path = convert_docx_to_pdf(docx_path)
        
        pdf_path_ans = ""
        if include_answers:
            docx_path_ans = create_ans_file(selected_words, start_no, end_no)
            if use_image_pdf:
                pdf_path_ans = convert_docx_to_image_pdf(docx_path_ans)
            else:
                pdf_path_ans = convert_docx_to_pdf(docx_path_ans)
        
        print_test(pdf_path, num_copies, include_answers, pdf_path_ans, printer_name)
    except ValueError as e:
        messagebox.showerror("入力エラー", str(e))
    except Exception as e:
        messagebox.showerror("エラー", f"予期しないエラーが発生しました:\n{e}")


# Wordファイルの作成
def create_word_file(selected_words, start_no, end_no):
    try:
        doc = docx.Document()

        # ページのマージンを設定
        sections = doc.sections
        for section in sections:
            section.top_margin = docx.shared.Cm(2.0)    # 上マージンを2cmに設定
            section.bottom_margin = docx.shared.Cm(2.0) # 下マージンを2cmに設定
            section.left_margin = docx.shared.Cm(2.0)   # 左マージンを2cmに設定
            section.right_margin = docx.shared.Cm(2.0)  # 右マージンを2cmに設定

        head = 'Vocabulary Test / ' + str(start_no) + ' ~ ' + str(end_no) + '       Name            Score'
        doc.add_paragraph(head).runs[0].font.size = Pt(18)

        for block_start in range(0, len(selected_words), 50):
            block = selected_words[block_start:block_start + 50]
            table = doc.add_table(rows=len(block)//2 + len(block)%2, cols=2)

            for i in range(0, len(block), 2):
                word1 = block[i]['Word']
                cell1 = table.cell(i//2, 0)
                paragraph1 = cell1.paragraphs[0]
                run1 = paragraph1.add_run(f"No.{block_start + i + 1} {word1}")
                run1.font.size = Pt(11)  # 文字サイズを9ptに変更
                run1.font.name = 'メイリオ'
                
                if i + 1 < len(block):
                    word2 = block[i + 1]['Word']
                    cell2 = table.cell(i//2, 1)
                    paragraph2 = cell2.paragraphs[0]
                    run2 = paragraph2.add_run(f"No.{block_start + i + 2} {word2}")
                    run2.font.size = Pt(11)  # 文字サイズを9ptに変更
                    run2.font.name = 'メイリオ'

                # 行の高さを調整（行のスペースを小さくする）
                paragraph_format1 = paragraph1.paragraph_format
                paragraph_format1.line_spacing = Pt(25)  # 行間を狭く設定
                paragraph_format1.space_after = Pt(0)  # 段落後のスペースを削除

                paragraph_format2 = paragraph2.paragraph_format
                paragraph_format2.line_spacing = Pt(10)
                paragraph_format2.space_after = Pt(0)

            # テーブルの罫線を非表示にする
            for row in table.rows:
                for cell in row.cells:
                    for border in ('top', 'bottom', 'left', 'right'):
                        cell._element.get_or_add_tcPr().append(docx.oxml.parse_xml(
                            f'<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                            f'<w:{border} w:val="none"/></w:tcBorders>'
                        ))

        desktop_path = os.path.join(os.path.join(os.environ['USERPROFILE']), 'Desktop')
        save_folder = os.path.join(desktop_path, '単語テスト')
        if not os.path.exists(save_folder):
            os.makedirs(save_folder)
        
        save_path = os.path.join(save_folder, 'vocabulary_test.docx')
        doc.save(save_path)
        print(f"File saved: {save_path}")
        return save_path
    except Exception as e:
        error_message = f"Wordファイルの作成中にエラーが発生しました:\n{str(e)}\n"
        error_message += f"Error type: {type(e).__name__}\n"
        messagebox.showerror("ファイル作成エラー", error_message)
        raise




# 答えのファイルを作成
def create_ans_file(selected_words, start_no, end_no):
    try:
        doc_ans = docx.Document()

        # ページのマージンを設定（各2.0cm）
        sections = doc_ans.sections
        for section in sections:
            section.top_margin = docx.shared.Cm(2.0)    # 上マージンを2cmに設定
            section.bottom_margin = docx.shared.Cm(2.0) # 下マージンを2cmに設定
            section.left_margin = docx.shared.Cm(2.0)   # 左マージンを2cmに設定
            section.right_margin = docx.shared.Cm(2.0)  # 右マージンを2cmに設定

        head_ans = 'Vocabulary Test Answers / ' + str(start_no) + ' ~ ' + str(end_no)
        doc_ans.add_paragraph(head_ans).runs[0].font.size = Pt(18)

        for block_start in range(0, len(selected_words), 50):
            block = selected_words[block_start:block_start + 50]
            table = doc_ans.add_table(rows=len(block)//2 + len(block)%2, cols=2)

            for i in range(0, len(block), 2):
                meaning1 = block[i]['meaning']
                cell1 = table.cell(i//2, 0)
                paragraph1 = cell1.paragraphs[0]
                run1 = paragraph1.add_run(f"No.{block_start + i + 1} {meaning1}")
                run1.font.size = Pt(11)  # 文字サイズを11ptに変更
                run1.font.name = 'メイリオ'

                if i + 1 < len(block):
                    meaning2 = block[i + 1]['meaning']
                    cell2 = table.cell(i//2, 1)
                    paragraph2 = cell2.paragraphs[0]
                    run2 = paragraph2.add_run(f"No.{block_start + i + 2} {meaning2}")
                    run2.font.size = Pt(11)  # 文字サイズを11ptに変更
                    run2.font.name = 'メイリオ'

                # 行の高さを調整
                paragraph_format1 = paragraph1.paragraph_format
                paragraph_format1.line_spacing = Pt(20)  # 行間を設定
                paragraph_format1.space_after = Pt(0)    # 段落後のスペースを削除

                paragraph_format2 = paragraph2.paragraph_format
                paragraph_format2.line_spacing = Pt(12)
                paragraph_format2.space_after = Pt(0)

            # テーブルの罫線を非表示にする
            for row in table.rows:
                for cell in row.cells:
                    for border in ('top', 'bottom', 'left', 'right'):
                        cell._element.get_or_add_tcPr().append(docx.oxml.parse_xml(
                            f'<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                            f'<w:{border} w:val="none"/></w:tcBorders>'
                        ))

        for paragraph in doc_ans.paragraphs:
            paragraph_format = paragraph.paragraph_format
            paragraph_format.space_after = Pt(0)

        desktop_path = os.path.join(os.path.join(os.environ['USERPROFILE']), 'Desktop')
        save_folder = os.path.join(desktop_path, '単語テスト')

        if not os.path.exists(save_folder):
            os.makedirs(save_folder)

        save_path_ans = os.path.join(save_folder, 'vocabulary_test_answers.docx')
        doc_ans.save(save_path_ans)

        print(f"File saved: {save_path_ans}")
        return save_path_ans

    except Exception as e:
        error_message = f"Wordファイルの作成中にエラーが発生しました:\n{str(e)}\n"
        error_message += f"Error type: {type(e).__name__}\n"
        messagebox.showerror("ファイル作成エラー", error_message)
        raise




# WordファイルをPDFに変換
def convert_docx_to_pdf(docx_path):
    try:
        pdf_path = docx_path.replace('.docx', '.pdf')
        doc = SimpleDocTemplate(pdf_path, pagesize=letter)
        story = []
        
        docx_file = docx.Document(docx_path)
        styles = getSampleStyleSheet()
        styleN = styles['Normal']
        
        for para in docx_file.paragraphs:
            if para.text.strip():
                story.append(Paragraph(para.text, styleN))
        
        for table in docx_file.tables:
            data = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                data.append(row_data)
            if data:
                pdf_table = Table(data)
                pdf_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), '#d5a3a3'),
                    ('TEXTCOLOR', (0, 0), (-1, 0), '#000000'),
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                    ('BACKGROUND', (0, 1), (-1, -1), '#f2f2f2'),
                    ('GRID', (0, 0), (-1, -1), 1, '#d5a3a3'),
                ]))
                story.append(pdf_table)
        
        doc.build(story)
        print(f"PDF file saved: {pdf_path}")
        return pdf_path
    except Exception as e:
        messagebox.showerror("PDF変換エラー", f"PDF変換中にエラーが発生しました\nもう一度やり直してください:\n{e}")
        raise


# Wordファイルを画像としてPDFに変換
poppler_path = os.path.join(os.getcwd(), 'poppler', 'bin')

def convert_docx_to_image_pdf(docx_path):
    pdf_path = docx_path.replace('.docx', '.pdf')
    convert(docx_path, pdf_path)

    # Popplerのパスを指定してPDFを画像に変換
    poppler_path = resource_path('poppler/bin')
    images = convert_from_path(pdf_path, poppler_path=poppler_path)

    image_pdf_path = pdf_path
    c = canvas.Canvas(image_pdf_path, pagesize=letter)
    
    # 画像を一時ファイルとして保存してからPDFに追加
    for image in images:
        with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as temp_file:
            image.save(temp_file.name, 'PNG')
            c.drawImage(temp_file.name, 0, 0, *letter)
            c.showPage()
        os.unlink(temp_file.name)  # 一時ファイルを削除
    
    c.save()
    print(f"PDF file saved as images: {image_pdf_path}")
    return image_pdf_path


# テストの印刷処理
def print_test(pdf_path, num_copies, include_answers, pdf_path_ans, printer_name):
    try:
        if printer_name not in printers:
            raise ValueError("選択されたプリンタが無効です。")
        
        for _ in range(num_copies):
            win32api.ShellExecute(0, "printto", pdf_path, f'"{printer_name}"', ".", 0)
            time.sleep(1)
        
        if include_answers:
            win32api.ShellExecute(0, "printto", pdf_path_ans, f'"{printer_name}"', ".", 0)
    except Exception as e:
        messagebox.showerror("印刷エラー", f"印刷中にエラーが発生しました\nやり直すか'{pdf_path}'を手動で印刷してください:\n{e}")
        raise


# GUIのセットアップ
root = tk.Tk()
root.title("単語テスト生成ツール")
printers = [printer[2] for printer in win32print.EnumPrinters(win32print.PRINTER_ENUM_LOCAL | win32print.PRINTER_ENUM_CONNECTIONS)]

printer_var = tk.StringVar(root)
printer_var.set(printers[0])

def set_printer():
    selected_printer = printer_var.get()
    win32print.SetDefaultPrinter(selected_printer)

book_label = tk.Label(root, text="単語帳を選択:")
book_label.grid(row=0, column=0, padx=10, pady=10, sticky="e")
book_var = tk.StringVar()
book_menu = ttk.Combobox(root, textvariable=book_var)
book_menu['values'] = ("システム英単語")
book_menu.grid(row=0, column=1, padx=10, pady=10)

start_no_label = tk.Label(root, text="開始No.:")
start_no_label.grid(row=1, column=0, padx=10, pady=10, sticky="e")
start_no_entry = tk.Entry(root)
start_no_entry.grid(row=1, column=1, padx=10, pady=10)

end_no_label = tk.Label(root, text="終了No.:")
end_no_label.grid(row=2, column=0, padx=10, pady=10, sticky="e")
end_no_entry = tk.Entry(root)
end_no_entry.grid(row=2, column=1, padx=10, pady=10)

num_words_label = tk.Label(root, text="テストする単語数:")
num_words_label.grid(row=3, column=0, padx=10, pady=10, sticky="e")
num_words_entry = tk.Entry(root)
num_words_entry.grid(row=3, column=1, padx=10, pady=10)

include_answers_var = tk.BooleanVar()
include_answers_check = tk.Checkbutton(root, text="答えも印刷", variable=include_answers_var)
include_answers_check.grid(row=4, column=0, columnspan=2, pady=10)

num_copies_label = tk.Label(root, text="印刷部数:")
num_copies_label.grid(row=5, column=0, padx=10, pady=10, sticky="e")
num_copies_entry = tk.Entry(root)
num_copies_entry.grid(row=5, column=1, padx=10, pady=10)

def on_generate_and_print():
    set_printer()
    create_test()

label = ttk.Label(root, text="プリンターを選択してください:")
label.grid(row=7, column=0, pady=10)

printer_menu = ttk.OptionMenu(root, printer_var, *printers)
printer_menu.grid(row=7, column=1, pady=10)

generate_button = tk.Button(root, text="印刷", command=on_generate_and_print)
generate_button.grid(row=8, column=0, columnspan=2, pady=10)

root.mainloop()