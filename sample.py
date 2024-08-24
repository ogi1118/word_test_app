import tkinter as tk
import docx
import os

def create_word_file():
    # ユーザーのデスクトップパスを取得
    desktop_path = os.path.join(os.path.join(os.environ['USERPROFILE']), 'Desktop')
    
    # "単語テスト"フォルダのパス
    save_folder = os.path.join(desktop_path, '単語テスト')
    
    # "単語テスト"フォルダが存在しない場合は作成
    if not os.path.exists(save_folder):
        os.makedirs(save_folder)
    
    # 保存ファイルのパス
    save_path = os.path.join(save_folder, 'sample.docx')
    
    # Wordファイルの作成と保存
    doc = docx.Document()
    doc.add_heading('Sample Document', 0)
    doc.add_paragraph('This is a sample paragraph.')
    doc.save(save_path)
    
    print(f"File saved: {save_path}")

# GUI setup
root = tk.Tk()
root.title("Sample App")

frame = tk.Frame(root)
frame.pack(padx=10, pady=10)

button = tk.Button(frame, text="Create Word File", command=create_word_file)
button.pack(pady=10)

root.mainloop()
